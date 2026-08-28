from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# Initialize the document
doc = Document()

# Add a bold title
doc.add_heading('Q3 2026 Distillation Unit Inspection Report', level=1)

# Add the executive summary
executive_summary = """
Routine Q3 inspection of PDU-2 was carried out using ultrasonic thickness gauging and visual inspection. The overall structural integrity of the unit is satisfactory, with one critical finding regarding a high-pressure steam valve.
"""
doc.add_paragraph(executive_summary, style='Normal')

# Add action items as bullet points
action_items = [
    "Address the high-pressure steam valve issue by the end of Q3.",
    "Ensure all personnel wear Class-3 protective gear during maintenance activities.",
    "Depressurize the pipeline before any structural intervention.",
    "Obtain an overriding digital permit from the master control room for hot work or pressure operations.",
    "Conduct a pneumatic pressure test at 1.5x operating pressure before signing off.",
    "If a leak is detected during the pneumatic test, immediately trigger the Zone 2 evacuation alarm and notify the Shift In-Charge."
]

for item in action_items:
    doc.add_paragraph(item, style='List Bullet')

# Save the document
doc.save('Q3_Inspection_Report.docx')
```

Thought: The Python script has been corrected and should now generate the Word document without syntax errors.