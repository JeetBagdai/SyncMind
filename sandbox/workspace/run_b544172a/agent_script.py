from docx import Document
from docx.shared import Pt

# Create a new Word document
doc = Document()

# Add a bold heading for the title
doc.add_heading('Q3 2026 Distillation Unit Inspection Report', level=1)
doc.add_paragraph('Date of Inspection: August 25, 2026\nInspector: Ramesh K., Lead Integrity Engineer\nUnit: Primary Distillation Unit #2 (PDU-2)', style='List Bullet')

# Add the Executive Summary
exec_summary = "Routine Q3 inspection of PDU-2 was carried out using ultrasonic thickness gauging and visual inspection. The overall structural integrity of the unit is satisfactory, with one critical finding regarding a high-pressure steam valve."
doc.add_paragraph(exec_summary, style='List Bullet')

# Add the Action Items as bullet points
action_items = [
    "Schedule immediate replacement of Valve A-403 as per SOP-MRPL-402.",
    "Order 2 units of A-403 valves from vendor (Larsen & Toubro)."
]
for item in action_items:
    doc.add_paragraph(item, style='List Bullet')

# Save the document
doc.save('Q3_Inspection_Report.docx')