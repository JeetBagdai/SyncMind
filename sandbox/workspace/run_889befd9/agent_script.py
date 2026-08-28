from docx import Document
from docx.shared import Pt

# Create a new Word document
doc = Document()

# Add the executive summary
doc.add_heading('Executive Summary', level=1)
executive_summary = "Routine Q3 inspection of PDU-2 was carried out using ultrasonic thickness gauging and visual inspection. The overall structural integrity of the unit is satisfactory, with one critical finding regarding a high-pressure steam valve."
doc.add_paragraph(executive_summary)

# Add the detailed findings
doc.add_heading('Detailed Findings', level=1)
findings = [
    "- Heat Exchangers (HX-101 to HX-104): No significant fouling detected. Thickness readings are within the 2mm corrosion allowance.",
    "- Piping Network (Line L-205): Minor surface oxidation observed on the external insulation cladding. Recommended for painting in Q4.",
    "- Valves and Instrumentation: Valve A-403 (HP Steam line) showed signs of micro-fissures on the upstream flange during dye-penetrant testing. The acoustic emission sensor indicated a faint internal bypass leak when fully closed."
]
for finding in findings:
    doc.add_paragraph(finding)

# Add the action items
doc.add_heading('Action Items', level=1)
action_items = [
    "1. Schedule immediate replacement of Valve A-403 as per SOP-MRPL-402.",
    "2. Order 2 units of A-403 valves from vendor (Larsen & Toubro)."
]
for item in action_items:
    doc.add_paragraph(item)

# Save the document
doc.save('Q3_Inspection_Report.docx')